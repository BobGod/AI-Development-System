#!/usr/bin/env python3
"""
智能行业知识问答系统 - 文档解析器
支持PDF、Word、Excel、图片等多种格式文档的智能解析
"""

import os
import io
import re
import json
from typing import Dict, List, Any, Optional, Union, Tuple
from pathlib import Path
from dataclasses import dataclass, asdict
from datetime import datetime
import asyncio
import logging

# 文档处理库
import pdfplumber  # 更好的PDF处理库
import fitz  # PyMuPDF，对中文支持更好
from docx import Document
import openpyxl

# OCR相关库 - 可选依赖
try:
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_path
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# NLP和内容分析 - 可选依赖
try:
    import jieba
    from langdetect import detect
    NLP_AVAILABLE = True
except ImportError:
    NLP_AVAILABLE = False

@dataclass
class DocumentMetadata:
    """文档元数据"""
    file_path: str
    file_name: str
    file_type: str
    file_size: int
    created_at: str
    modified_at: str
    page_count: int = 0
    author: str = ""
    title: str = ""
    language: str = ""
    encoding: str = "utf-8"

@dataclass
class ExtractedContent:
    """提取的内容"""
    content_id: str
    content_type: str  # text, table, image, list, etc.
    content: str
    page_number: int = 0
    position: Dict[str, Any] = None
    confidence: float = 1.0
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.position is None:
            self.position = {}
        if self.metadata is None:
            self.metadata = {}

@dataclass 
class DocumentParseResult:
    """文档解析结果"""
    document_id: str
    metadata: DocumentMetadata
    extracted_contents: List[ExtractedContent]
    summary: str = ""
    keywords: List[str] = None
    topics: List[str] = None
    parse_status: str = "success"
    error_message: str = ""
    parsed_at: str = ""
    
    def __post_init__(self):
        if self.keywords is None:
            self.keywords = []
        if self.topics is None:
            self.topics = []
        if not self.parsed_at:
            self.parsed_at = datetime.now().isoformat()

class DocumentParser:
    """文档解析器主类"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.logger = logging.getLogger(__name__)
        
        # 支持的文件类型
        self.supported_extensions = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.xlsx': self._parse_xlsx,
            '.xls': self._parse_xls,
            '.txt': self._parse_txt,
            '.md': self._parse_markdown,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image,
            '.png': self._parse_image,
            '.bmp': self._parse_image,
            '.tiff': self._parse_image
        }
        
        # OCR配置
        self.ocr_config = self.config.get('ocr', {
            'lang': 'chi_sim+eng',  # 中英文
            'psm': 6,  # 页面分割模式
            'oem': 3   # OCR引擎模式
        })
        
        # 内容清理正则表达式
        self.cleanup_patterns = [
            r'\s+',  # 多个空白字符
            r'\n\s*\n',  # 多个换行
            r'[^\w\s\u4e00-\u9fff]+'  # 非中英文字符和数字
        ]
        
    async def parse_document(self, file_path: str) -> DocumentParseResult:
        """解析单个文档"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
                
            # 获取文件扩展名
            extension = file_path.suffix.lower()
            if extension not in self.supported_extensions:
                raise ValueError(f"不支持的文件类型: {extension}")
                
            self.logger.info(f"开始解析文档: {file_path}")
            
            # 获取文档元数据
            metadata = self._extract_metadata(file_path)
            
            # 根据文件类型选择解析器
            parser_func = self.supported_extensions[extension]
            extracted_contents = await parser_func(file_path, metadata)
            
            # 生成文档摘要和关键词
            all_text = "\n".join([content.content for content in extracted_contents 
                                 if content.content_type == 'text'])
            
            summary = self._generate_summary(all_text)
            keywords = self._extract_keywords(all_text)
            topics = self._extract_topics(all_text)
            
            # 创建解析结果
            result = DocumentParseResult(
                document_id=self._generate_document_id(file_path),
                metadata=metadata,
                extracted_contents=extracted_contents,
                summary=summary,
                keywords=keywords,
                topics=topics
            )
            
            self.logger.info(f"文档解析完成: {file_path}, 提取内容 {len(extracted_contents)} 个")
            return result
            
        except Exception as e:
            self.logger.error(f"文档解析失败: {file_path}, 错误: {e}")
            return DocumentParseResult(
                document_id=self._generate_document_id(file_path),
                metadata=self._extract_metadata(file_path),
                extracted_contents=[],
                parse_status="failed",
                error_message=str(e)
            )
            
    async def parse_batch(self, file_paths: List[str], 
                         max_concurrent: int = 5) -> List[DocumentParseResult]:
        """批量解析文档"""
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def parse_with_semaphore(file_path: str) -> DocumentParseResult:
            async with semaphore:
                return await self.parse_document(file_path)
                
        tasks = [parse_with_semaphore(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 处理异常结果
        parsed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                self.logger.error(f"批量解析失败: {file_paths[i]}, 错误: {result}")
                parsed_results.append(DocumentParseResult(
                    document_id=self._generate_document_id(Path(file_paths[i])),
                    metadata=self._extract_metadata(Path(file_paths[i])),
                    extracted_contents=[],
                    parse_status="failed",
                    error_message=str(result)
                ))
            else:
                parsed_results.append(result)
                
        return parsed_results
        
    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """提取文档元数据"""
        try:
            stat = file_path.stat()
            
            metadata = DocumentMetadata(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=file_path.suffix.lower(),
                file_size=stat.st_size,
                created_at=datetime.fromtimestamp(stat.st_ctime).isoformat(),
                modified_at=datetime.fromtimestamp(stat.st_mtime).isoformat()
            )
            
            return metadata
            
        except Exception as e:
            self.logger.warning(f"提取元数据失败: {file_path}, 错误: {e}")
            return DocumentMetadata(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=file_path.suffix.lower(),
                file_size=0,
                created_at=datetime.now().isoformat(),
                modified_at=datetime.now().isoformat()
            )
            
    async def _parse_pdf(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析PDF文档"""
        contents = []
        content_counter = 0
        
        try:
            # 优先使用pdfplumber，对中文支持更好
            self.logger.info(f"使用pdfplumber解析PDF: {file_path}")
            with pdfplumber.open(str(file_path)) as pdf:
                metadata.page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            # pdfplumber对中文的处理更好
                            clean_text = self._clean_text(text)
                            if clean_text:
                                contents.append(ExtractedContent(
                                    content_id=f"pdf_text_{content_counter}",
                                    content_type="text",
                                    content=clean_text,
                                    page_number=page_num,
                                    confidence=0.95,  # pdfplumber置信度更高
                                    metadata={"extraction_method": "pdfplumber"}
                                ))
                                content_counter += 1
                    except Exception as e:
                        self.logger.warning(f"pdfplumber页面 {page_num} 文本提取失败: {e}")
                        
            # 如果pdfplumber提取效果不好，尝试PyMuPDF
            if len(contents) == 0 or sum(len(c.content) for c in contents) < 100:
                self.logger.info(f"pdfplumber提取不足，尝试PyMuPDF: {file_path}")
                contents.extend(await self._parse_pdf_with_pymupdf(file_path, metadata))
                
            # 如果还是不够，使用OCR作为最后手段
            if len(contents) == 0 or sum(len(c.content) for c in contents) < 50:
                self.logger.info(f"文本提取仍不足，使用OCR: {file_path}")
                ocr_contents = await self._parse_pdf_with_ocr(file_path, metadata)
                contents.extend(ocr_contents)
                
        except Exception as e:
            self.logger.error(f"PDF解析失败: {file_path}, 错误: {e}")
            # 尝试备选方案
            try:
                contents = await self._parse_pdf_with_pymupdf(file_path, metadata)
                if not contents:
                    contents = await self._parse_pdf_with_ocr(file_path, metadata)
            except Exception as fallback_error:
                self.logger.error(f"所有PDF解析方法都失败: {fallback_error}")
                
        return contents
        
    async def _parse_pdf_with_pymupdf(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """使用PyMuPDF解析PDF"""
        contents = []
        content_counter = 0
        
        try:
            self.logger.info(f"使用PyMuPDF解析PDF: {file_path}")
            # 打开PDF文档
            doc = fitz.open(str(file_path))
            metadata.page_count = len(doc)
            
            for page_num in range(len(doc)):
                try:
                    page = doc[page_num]
                    # 提取文本，PyMuPDF对中文支持较好
                    text = page.get_text()
                    
                    if text and text.strip():
                        clean_text = self._clean_text(text)
                        if clean_text:
                            contents.append(ExtractedContent(
                                content_id=f"pdf_pymupdf_{content_counter}",
                                content_type="text",
                                content=clean_text,
                                page_number=page_num + 1,
                                confidence=0.9,
                                metadata={"extraction_method": "pymupdf"}
                            ))
                            content_counter += 1
                            
                except Exception as e:
                    self.logger.warning(f"PyMuPDF页面 {page_num + 1} 文本提取失败: {e}")
                    
            doc.close()
            
        except Exception as e:
            self.logger.error(f"PyMuPDF解析失败: {e}")
            
        return contents
        
    async def _parse_pdf_with_ocr(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """使用OCR解析PDF"""
        contents = []
        content_counter = 0
        
        try:
            # 将PDF转换为图像
            images = convert_from_path(str(file_path))
            metadata.page_count = len(images)
            
            for page_num, image in enumerate(images, 1):
                try:
                    # OCR识别
                    text = pytesseract.image_to_string(
                        image, 
                        lang=self.ocr_config['lang'],
                        config=f'--psm {self.ocr_config["psm"]} --oem {self.ocr_config["oem"]}'
                    )
                    
                    if text.strip():
                        contents.append(ExtractedContent(
                            content_id=f"pdf_ocr_{content_counter}",
                            content_type="text",
                            content=self._clean_text(text),
                            page_number=page_num,
                            confidence=0.7,  # OCR置信度通常较低
                            metadata={"extraction_method": "ocr"}
                        ))
                        content_counter += 1
                        
                except Exception as e:
                    self.logger.warning(f"PDF页面 {page_num} OCR失败: {e}")
                    
        except Exception as e:
            self.logger.error(f"PDF OCR处理失败: {e}")
            
        return contents
        
    async def _parse_docx(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析Word文档(.docx)"""
        contents = []
        content_counter = 0
        
        try:
            doc = Document(str(file_path))
            
            # 提取文档属性
            if doc.core_properties.title:
                metadata.title = doc.core_properties.title
            if doc.core_properties.author:
                metadata.author = doc.core_properties.author
                
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    contents.append(ExtractedContent(
                        content_id=f"docx_para_{content_counter}",
                        content_type="text",
                        content=self._clean_text(para.text),
                        confidence=1.0,
                        metadata={"style": para.style.name if para.style else "Normal"}
                    ))
                    content_counter += 1
                    
            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    
                if table_data:
                    # 将表格转换为文本格式
                    table_text = self._table_to_text(table_data)
                    contents.append(ExtractedContent(
                        content_id=f"docx_table_{table_idx}",
                        content_type="table",
                        content=table_text,
                        confidence=1.0,
                        metadata={"table_data": table_data}
                    ))
                    
        except Exception as e:
            self.logger.error(f"DOCX解析失败: {file_path}, 错误: {e}")
            
        return contents
        
    async def _parse_doc(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析Word文档(.doc) - 需要转换或使用其他库"""
        # 简化实现：建议用户转换为.docx格式
        contents = []
        contents.append(ExtractedContent(
            content_id="doc_notice",
            content_type="text",
            content=f"检测到.doc格式文档：{file_path.name}\n建议转换为.docx格式以获得更好的解析效果。",
            confidence=0.5,
            metadata={"notice": "format_conversion_needed"}
        ))
        return contents
        
    async def _parse_xlsx(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析Excel文档(.xlsx)"""
        contents = []
        content_counter = 0
        
        try:
            workbook = openpyxl.load_workbook(str(file_path), data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # 获取有数据的区域
                if sheet.max_row > 0 and sheet.max_column > 0:
                    sheet_data = []
                    
                    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row,
                                             min_col=1, max_col=sheet.max_column,
                                             values_only=True):
                        # 过滤空行
                        if any(cell is not None and str(cell).strip() for cell in row):
                            sheet_data.append([str(cell) if cell is not None else "" for cell in row])
                            
                    if sheet_data:
                        # 将表格转换为文本
                        table_text = self._table_to_text(sheet_data)
                        contents.append(ExtractedContent(
                            content_id=f"xlsx_sheet_{content_counter}",
                            content_type="table",
                            content=table_text,
                            confidence=1.0,
                            metadata={
                                "sheet_name": sheet_name,
                                "table_data": sheet_data,
                                "rows": len(sheet_data),
                                "columns": len(sheet_data[0]) if sheet_data else 0
                            }
                        ))
                        content_counter += 1
                        
        except Exception as e:
            self.logger.error(f"XLSX解析失败: {file_path}, 错误: {e}")
            
        return contents
        
    async def _parse_xls(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析Excel文档(.xls) - 建议转换为.xlsx"""
        contents = []
        contents.append(ExtractedContent(
            content_id="xls_notice",
            content_type="text",
            content=f"检测到.xls格式文档：{file_path.name}\n建议转换为.xlsx格式以获得更好的解析效果。",
            confidence=0.5,
            metadata={"notice": "format_conversion_needed"}
        ))
        return contents
        
    async def _parse_txt(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析文本文档"""
        contents = []
        
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            text_content = None
            used_encoding = 'utf-8'
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
                    
            if text_content:
                metadata.encoding = used_encoding
                # 检测语言
                try:
                    metadata.language = detect(text_content[:1000])  # 用前1000字符检测语言
                except:
                    metadata.language = 'unknown'
                    
                contents.append(ExtractedContent(
                    content_id="txt_content",
                    content_type="text", 
                    content=self._clean_text(text_content),
                    confidence=1.0,
                    metadata={"encoding": used_encoding}
                ))
                
        except Exception as e:
            self.logger.error(f"TXT解析失败: {file_path}, 错误: {e}")
            
        return contents
        
    async def _parse_markdown(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析Markdown文档"""
        contents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                markdown_content = file.read()
                
            # 检测语言
            try:
                metadata.language = detect(markdown_content[:1000])
            except:
                metadata.language = 'unknown'
                
            contents.append(ExtractedContent(
                content_id="md_content",
                content_type="text",
                content=self._clean_text(markdown_content),
                confidence=1.0,
                metadata={"format": "markdown"}
            ))
            
        except Exception as e:
            self.logger.error(f"Markdown解析失败: {file_path}, 错误: {e}")
            
        return contents
        
    async def _parse_image(self, file_path: Path, metadata: DocumentMetadata) -> List[ExtractedContent]:
        """解析图片文档(OCR)"""
        contents = []
        
        try:
            # 读取图片
            image = Image.open(str(file_path))
            
            # 图片预处理（可选）
            image = self._preprocess_image(image)
            
            # OCR识别
            text = pytesseract.image_to_string(
                image,
                lang=self.ocr_config['lang'],
                config=f'--psm {self.ocr_config["psm"]} --oem {self.ocr_config["oem"]}'
            )
            
            if text.strip():
                contents.append(ExtractedContent(
                    content_id="image_ocr",
                    content_type="text",
                    content=self._clean_text(text),
                    confidence=0.7,  # OCR置信度
                    metadata={
                        "extraction_method": "ocr",
                        "image_size": image.size,
                        "image_mode": image.mode
                    }
                ))
                
        except Exception as e:
            self.logger.error(f"图片解析失败: {file_path}, 错误: {e}")
            
        return contents
        
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """图片预处理，提高OCR准确率"""
        try:
            # 转换为OpenCV格式
            img_array = np.array(image)
            
            # 转换为灰度图
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
                
            # 二值化
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # 降噪
            denoised = cv2.medianBlur(binary, 3)
            
            # 转换回PIL格式
            processed_image = Image.fromarray(denoised)
            return processed_image
            
        except Exception as e:
            self.logger.warning(f"图片预处理失败: {e}")
            return image
            
    def _clean_text(self, text: str) -> str:
        """清理文本内容"""
        if not text:
            return ""
            
        # 去除多余空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 去除多余换行
        text = re.sub(r'\n\s*\n', '\n', text)
        
        # 去除首尾空白
        text = text.strip()
        
        return text
        
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """将表格数据转换为文本格式"""
        if not table_data:
            return ""
            
        text_lines = []
        for row in table_data:
            # 过滤空单元格
            filtered_row = [cell.strip() for cell in row if cell.strip()]
            if filtered_row:
                text_lines.append(" | ".join(filtered_row))
                
        return "\n".join(text_lines)
        
    def _generate_summary(self, text: str, max_length: int = 200) -> str:
        """生成文档摘要"""
        if not text or len(text) < 100:
            return text[:max_length] if text else ""
            
        # 简单的摘要生成：取前几句话
        sentences = re.split(r'[。！？\n]', text)
        summary_parts = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and current_length + len(sentence) < max_length:
                summary_parts.append(sentence)
                current_length += len(sentence)
            else:
                break
                
        summary = "。".join(summary_parts)
        if summary and not summary.endswith('。'):
            summary += "。"
            
        return summary if summary else text[:max_length]
        
    def _extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """提取关键词"""
        if not text:
            return []
            
        try:
            # 使用jieba进行中文分词
            words = jieba.cut(text)
            
            # 过滤停用词和短词
            stop_words = {'的', '是', '在', '了', '和', '有', '我', '你', '他', '她', '它', 
                         'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to'}
            
            filtered_words = []
            for word in words:
                word = word.strip()
                if (len(word) >= 2 and 
                    word not in stop_words and 
                    not word.isdigit() and
                    re.match(r'^[\w\u4e00-\u9fff]+$', word)):
                    filtered_words.append(word)
                    
            # 统计词频
            word_count = {}
            for word in filtered_words:
                word_count[word] = word_count.get(word, 0) + 1
                
            # 按频率排序并返回前N个
            sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)
            keywords = [word for word, count in sorted_words[:max_keywords]]
            
            return keywords
            
        except Exception as e:
            self.logger.warning(f"关键词提取失败: {e}")
            return []
            
    def _extract_topics(self, text: str) -> List[str]:
        """提取主题"""
        if not text:
            return []
            
        # 简单的主题提取：基于关键词聚类
        keywords = self._extract_keywords(text, 20)
        
        # 这里可以实现更复杂的主题建模算法
        # 目前返回前几个关键词作为主题
        return keywords[:5]
        
    def _generate_document_id(self, file_path: Path) -> str:
        """生成文档唯一ID"""
        import hashlib
        
        # 使用文件路径和修改时间生成唯一ID
        content = f"{file_path.absolute()}_{file_path.stat().st_mtime}"
        return hashlib.md5(content.encode()).hexdigest()
        
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        return list(self.supported_extensions.keys())
        
    def is_supported_format(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_extensions

# 使用示例
async def main():
    """测试文档解析器"""
    parser = DocumentParser()
    
    # 测试单个文档解析
    test_file = "test_document.pdf"  # 替换为实际文件路径
    if Path(test_file).exists():
        result = await parser.parse_document(test_file)
        print(f"解析结果: {result.parse_status}")
        print(f"提取内容数量: {len(result.extracted_contents)}")
        print(f"摘要: {result.summary}")
        print(f"关键词: {result.keywords}")
        
if __name__ == "__main__":
    asyncio.run(main())